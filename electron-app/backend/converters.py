"""
This module contains all the conversion and processing functions
adapted from the original Python scripts.
"""

import os
import re
import subprocess
import shutil
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn
from docx.shared import Inches
from pdf2docx import Converter
import tempfile
import pandas as pd
import sys

# Helper function to get resource paths (works in both dev and PyInstaller)
def get_resource_path(relative_path):
    """
    Get absolute path to resource, works for dev and for PyInstaller
    """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        # Not running as PyInstaller executable
        base_path = os.path.dirname(os.path.abspath(__file__))
    
    # Check if we're in a PyInstaller onedir build (_internal folder)
    internal_path = os.path.join(base_path, '_internal', relative_path)
    if os.path.exists(internal_path):
        return internal_path
    
    # Standard path (development or direct in base)
    standard_path = os.path.join(base_path, relative_path)
    return standard_path

def get_tool_path(tool_name):
    """
    Get the full path to a bundled tool executable.
    
    Args:
        tool_name: Name of the tool (e.g., 'pandoc', 'asciidoctor')
    
    Returns:
        Full path to the executable, or just the tool name if not found (fallback to PATH)
    """
    if tool_name == 'pandoc':
        pandoc_path = get_resource_path(os.path.join('tools', 'pandoc', 'pandoc.exe'))
        if os.path.exists(pandoc_path):
            return pandoc_path
    elif tool_name == 'asciidoctor':
        # Use bundled Ruby runtime to run asciidoctor
        asciidoctor_bat = get_resource_path(os.path.join('ruby-runtime', 'bin', 'asciidoctor.bat'))
        if os.path.exists(asciidoctor_bat):
            return asciidoctor_bat
        # Fallback: try without .bat extension
        asciidoctor_path = get_resource_path(os.path.join('ruby-runtime', 'bin', 'asciidoctor'))
        if os.path.exists(asciidoctor_path):
            return asciidoctor_path
    
    # Fallback to tool name (will use PATH)
    return tool_name

# ============================================================================
# 1. DOCX to S1000D AsciiDoc Converter
# ============================================================================

def cleanup_adoc_content(content):
    """Apply cleanup transformations to AsciiDoc content."""
    # FIX 1: Remove trailing '+' at end of a line
    content = re.sub(r'(.+?)\s*\+\s*$', r'\1', content, flags=re.MULTILINE)
    
    # FIX 2: Replace standalone "{plus}" with "+"
    content = re.sub(r'^\s*\{plus\}\s*$', '+', content, flags=re.MULTILINE)
    
    # FIX 3: Remove blank lines inside fault blocks (`--`)
    content = re.sub(r'(--\n)\s+', r'\1', content)
    content = re.sub(r'\s+(\n--)', r'\1', content)
    
    # FIX 4: Join multiline attribute blocks `[ ... ]`
    lines = content.split('\n')
    rebuilt = []
    in_block = False
    
    for line in lines:
        trim = line.strip()
        if trim.startswith('[') and not trim.endswith(']'):
            in_block = True
            rebuilt.append(line.rstrip())
        elif in_block:
            rebuilt[-1] += ' ' + trim.replace('`', '')
            if trim.endswith(']'):
                in_block = False
        else:
            rebuilt.append(line)
    
    content = '\n'.join(rebuilt)
    
    # FIX 5: Normalize thematic breaks (---)
    content = re.sub(r'\s*^\s*-{3,}\s*$\s*', '\n\n---\n\n', content, flags=re.MULTILINE)
    
    return content

def convert_to_11_part_dmc(base_code):
    """Convert a 9-part DMC to 11-part format."""
    parts = base_code.split('-')
    
    if len(parts) == 9:
        p1, p2, p3, p4, p5, p6, p7, p8, p9 = parts
        sub_system_code = p4[0] if len(p4) > 0 else ''
        sub_sub_system_code = p4[1] if len(p4) > 1 else ''
        disassy_code = p7[:-1] if len(p7) > 0 else ''
        disassy_code_variant = p7[-1] if len(p7) > 0 else ''
        info_code = p8[:-1] if len(p8) > 0 else ''
        info_code_variant = p8[-1] if len(p8) > 0 else ''
        final_dmc = f"{p1}-{p2}-{p3}-{sub_system_code}-{sub_sub_system_code}-{p5}-{p6}-{disassy_code}-{disassy_code_variant}-{info_code}-{info_code_variant}-{p9}"
        return final_dmc, True
    
    return base_code, False

def create_procedural_header(dmc):
    """Create the S1000D metadata header for a Procedural AsciiDoc file."""
    header = f"""= My Procedural Data Module
:dmc: DMC-{dmc}
:dm-type: procedural
:issue-number: 001
:issue-date: 2023-10-26
:tech-name: Comprehensive Converter Test Procedure
:dm-title: Step-by-Step Guide
:revdate: 2025-09-02
:in-work: 00
:lang: en
:country-code: IN
:security-classification: 01
:responsible-partner-company: LNTDEFENCE
:enterprise-code-rpc: 1671Y
:originator-enterprise: LNTDEFENCE
:enterprise-code-originator: 1671Y
:applicability: All applicable units and serial numbers.
:brex-dmc: DMC-GSV-H-041-1-0-0301-00-A-022-A-D
:reason-for-update: Initial draft for demonstration purposes.
:s1000d-schema-base-path: http://www.s1000d.org/S1000D_4-2/xml_schema_flat/

[[prelim_reqs]]
== Preliminary Requirements

[[required_conditions_pr]]
=== Required Conditions

[[required_persons_pr]]
=== Required Persons

[[required_tech_info_pr]]
=== Required Technical Information

[[required_equip_pr]]
=== Required Support Equipment

[[required_supplies_pr]]
=== Required Supplies

[[required_spares_pr]]
=== Required Spares

[[required_safety_pr]]
=== Required Safety

[[main_proc_steps]]
== Main Procedure

"""
    return header

def create_descriptive_header(dmc):
    """Create the S1000D metadata header for a Descriptive AsciiDoc file."""
    header = f""":dmc: DMC-{dmc}
:dm-type: descript
:issue-number: 001
:dm-title: Sample Descriptive Module
:revdate: 2025-09-02
:in-work: 00
:lang: en
:country-code: IN
:security-classification: 01
:responsible-partner-company: LNTDEFENCE
:enterprise-code-rpc: 1671Y
:originator-enterprise: LNTDEFENCE
:enterprise-code-originator: 1671Y
:applicability: All applicable units and serial numbers.
:brex-dmc: DMC-GSV-H-041-1-0-0301-00-A-022-A-D
:reason-for-update: Initial draft for demonstration purposes.

"""
    return header

def create_procedural_footer():
    """Create the S1000D footer for a Procedural AsciiDoc file."""
    footer = """

[[closeout_reqs]]
== Closeout Requirements

[[closeout_conds_after]]
=== Required Conditions After Job Completion

"""
    return footer

def create_asciidoc_header(dmc, doc_type='descript'):
    """Create the S1000D metadata header for an AsciiDoc file.
    
    Args:
        dmc: The Data Module Code
        doc_type: 'proced' for procedural, 'descript' for descriptive (default)
    """
    if doc_type == 'proced':
        return create_procedural_header(dmc)
    else:
        return create_descriptive_header(dmc)

def determine_doc_type_from_dmc(base_code, was_converted):
    """Determine document type (proced/descript) based on info code in DMC.
    
    If info code is '000', it's procedural. Otherwise, it's descriptive.
    """
    if not was_converted:
        return 'descript'
    
    parts = base_code.split('-')
    if len(parts) >= 8:
        # In 9-part DMC, info code is in position 8 (index 7)
        info_code_part = parts[7] if len(parts) > 7 else ''
        # Extract info code (first 3 chars, excluding variant)
        info_code = info_code_part[:-1] if len(info_code_part) > 0 else ''
        if info_code == '000':
            return 'proced'
    return 'descript'

def convert_docx_to_s1000d(input_path, output_path, doc_type=None):
    """Convert a DOCX file to S1000D AsciiDoc format.
    
    Args:
        input_path: Path to the input DOCX file
        output_path: Path for the output ADOC file
        doc_type: 'proced' for procedural, 'descript' for descriptive, 
                  None for auto-detect based on DMC info code
    """
    try:
        # Get full path to pandoc executable
        pandoc_exe = get_tool_path('pandoc')
        
        # Run Pandoc conversion with explicit UTF-8 encoding
        result = subprocess.run(
            [pandoc_exe, str(input_path), '-t', 'asciidoc'],
            capture_output=True,
            text=True,
            encoding='utf-8',
            errors='replace',  # Replace characters that can't be decoded
            check=True
        )
        content = result.stdout
        
        # Get DMC from filename
        filename = os.path.splitext(os.path.basename(input_path))[0]
        base_code = filename.replace('DMC-', '', 1)
        final_dmc, was_converted = convert_to_11_part_dmc(base_code)
        
        # Determine document type
        if doc_type is None:
            # Auto-detect based on DMC info code
            doc_type = determine_doc_type_from_dmc(base_code, was_converted)
        
        # Create header and clean content
        header = create_asciidoc_header(final_dmc, doc_type)
        cleaned_content = cleanup_adoc_content(content)
        
        # Build final content
        if doc_type == 'proced':
            final_content = header + cleaned_content + create_procedural_footer()
        else:
            final_content = header + cleaned_content
        
        # Write output with UTF-8 encoding
        with open(output_path, 'w', encoding='utf-8', errors='replace') as f:
            f.write(final_content)
        
        return True, f"Conversion successful (type: {doc_type})"
    except subprocess.CalledProcessError as e:
        return False, f"Pandoc conversion failed: {e}"
    except FileNotFoundError:
        return False, f"Pandoc executable not found at: {pandoc_exe}"
    except Exception as e:
        return False, str(e)

# ============================================================================
# 2. PDF to DOCX Converter
# ============================================================================

def convert_pdf_to_docx(input_path, output_path):
    """Convert a PDF file to DOCX format."""
    try:
        converter = Converter(input_path)
        converter.convert(output_path)
        converter.close()
        return True, "Conversion successful"
    except Exception as e:
        return False, str(e)

# ============================================================================
# 3. DOCX Splitter by Heading
# ============================================================================

# Namespace map including 'wp' for drawing properties (like image size)
NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}

def iter_block_items(parent):
    """
    Yields each paragraph and table child within a parent element.
    Works for Document, _Cell, etc.
    """
    if hasattr(parent, "element") and hasattr(parent.element, "body"):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._element
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)

def copy_numbering(source_para, target_para):
    """
    Copies numbering (w:numPr) from source paragraph to target paragraph safely.
    """
    try:
        source_pPr = source_para._p.pPr
        if source_pPr is None:
            return
        numPr = source_pPr.find(qn("w:numPr"))
        if numPr is not None:
            target_pPr = target_para._p.get_or_add_pPr()
            existing_numPr = target_pPr.find(qn("w:numPr"))
            if existing_numPr is not None:
                target_pPr.remove(existing_numPr)
            target_pPr.append(numPr)
    except Exception:
        traceback.print_exc()

def safe_set_run_font(target_run, source_run):
    """
    Safely copy run-level font attributes (name, size, color) where possible.
    """
    try:
        if source_run.bold is not None:
            target_run.bold = source_run.bold
    except Exception:
        pass
    try:
        if source_run.italic is not None:
            target_run.italic = source_run.italic
    except Exception:
        pass
    try:
        if source_run.underline is not None:
            target_run.underline = source_run.underline
    except Exception:
        pass

    try:
        if hasattr(source_run, "font") and source_run.font is not None:
            if getattr(source_run.font, "name", None):
                target_run.font.name = source_run.font.name
            if getattr(source_run.font, "size", None):
                target_run.font.size = source_run.font.size

            try:
                color = getattr(source_run.font, "color", None)
                if color is not None:
                    rgb_val = getattr(color, "rgb", None)
                    if rgb_val is not None:
                        target_run.font.color.rgb = rgb_val
            except Exception:
                pass
    except Exception:
        traceback.print_exc()

def copy_paragraph(source_para, target_container, temp_img_dir):
    """
    Copy paragraph text, formatting and images from source_para into target_container
    (which can be a Document or a _Cell).
    """
    try:
        has_drawings = bool(source_para._p.findall(".//w:drawing", namespaces=NSMAP))
    except Exception:
        has_drawings = False

    if not source_para.text.strip() and not has_drawings:
        target_container.add_paragraph()
        return

    target_para = target_container.add_paragraph()

    try:
        if getattr(source_para, "style", None):
            target_para.style = source_para.style
    except Exception:
        pass

    try:
        target_para.paragraph_format.alignment = source_para.paragraph_format.alignment
    except Exception:
        pass

    try:
        copy_numbering(source_para, target_para)
    except Exception:
        pass

    for run in source_para.runs:
        try:
            target_run = target_para.add_run(run.text)
            safe_set_run_font(target_run, run)
        except Exception:
            try:
                target_run = target_para.add_run(run.text)
            except Exception:
                continue

        try:
            drawings = run._element.findall(".//w:drawing", namespaces=NSMAP)
        except Exception:
            drawings = []

        for drawing in drawings:
            try:
                for blip in drawing.findall(".//a:blip", namespaces=NSMAP):
                    rId = blip.get(qn("r:embed"))
                    if not rId:
                        continue
                    
                    image_part = None
                    try:
                        image_part = source_para.part.related_parts.get(rId)
                        if image_part is None:
                            image_part = source_para.part.related_parts[rId]
                    except Exception:
                        image_part = None

                    if image_part is None:
                        continue

                    try:
                        image_bytes = image_part.blob
                    except Exception:
                        continue

                    extension = "png"
                    try:
                        if hasattr(image_part, "ext") and image_part.ext:
                            extension = image_part.ext
                        else:
                            ext_candidate = getattr(image_part, "content_type", "")
                            if "/" in ext_candidate:
                                extension = ext_candidate.split("/")[-1]
                    except Exception:
                        pass
                    if extension == "jpeg":
                        extension = "jpg"

                    img_filename = f"{rId}.{extension}"
                    img_path = os.path.join(temp_img_dir, img_filename)
                    try:
                        if not os.path.exists(img_path):
                            with open(img_path, "wb") as f:
                                f.write(image_bytes)
                    except Exception:
                        continue

                    width = Inches(4)
                    try:
                        extent_el = drawing.find(".//wp:extent", namespaces=NSMAP)
                        if extent_el is not None:
                            cx = extent_el.get("cx")
                            if cx:
                                width = Inches(int(cx) / 914400)
                    except Exception:
                        pass

                    try:
                        pic_run = target_para.add_run()
                        pic_run.add_picture(img_path, width=width)
                    except Exception:
                        try:
                            container_doc = (
                                target_container
                                if isinstance(target_container, Document)
                                else target_container.part.package
                            )
                        except Exception:
                            container_doc = None
                        continue
            except Exception:
                traceback.print_exc()
                continue

def copy_table(source_table, target_container, temp_img_dir):
    """
    Copy table structure and content from source_table into target_container (Document or _Cell).
    Handles nested tables by rendering them as paragraphs inside the target cell.
    """
    try:
        if isinstance(target_container, type(Document())):
            target_table = target_container.add_table(rows=0, cols=len(source_table.columns))
        else:
            for source_row in source_table.rows:
                for source_cell in source_row.cells:
                    for block in iter_block_items(source_cell):
                        if isinstance(block, Paragraph):
                            copy_paragraph(block, target_container, temp_img_dir)
                        elif isinstance(block, Table):
                            copy_table(block, target_container, temp_img_dir)
            return

        try:
            target_table.style = source_table.style
        except Exception:
            pass
        try:
            target_table.autofit = source_table.autofit
        except Exception:
            pass

        for source_row in source_table.rows:
            target_row_cells = target_table.add_row().cells
            for i, source_cell in enumerate(source_row.cells):
                if i >= len(target_row_cells):
                    target_cell = target_row_cells[-1]
                else:
                    target_cell = target_row_cells[i]

                try:
                    target_cell._element.clear_content()
                except Exception:
                    pass

                for block in iter_block_items(source_cell):
                    if isinstance(block, Paragraph):
                        copy_paragraph(block, target_cell, temp_img_dir)
                    elif isinstance(block, Table):
                        copy_table(block, target_cell, temp_img_dir)
    except Exception:
        traceback.print_exc()

def split_docx_by_heading_v2(input_path, output_dir, heading_style='Heading 1'):
    """
    Enhanced version: Split a DOCX file into multiple files based on heading style.
    Advanced handling of images, tables (including nested), numbering, and formatting.
    """
    os.makedirs(output_dir, exist_ok=True)
    temp_img_dir = os.path.join(output_dir, "temp_images")
    os.makedirs(temp_img_dir, exist_ok=True)

    try:
        doc = Document(input_path)
        blocks = list(iter_block_items(doc))

        if not blocks:
            return 0

        # Determine first heading
        first = blocks[0]
        first_is_heading = (
            isinstance(first, Paragraph)
            and first.style is not None
            and getattr(first.style, "name", None) == heading_style
        )

        section_indices = []
        if not first_is_heading:
            section_indices.append(0)

        # Identify sections
        for i, block in enumerate(blocks):
            if (
                isinstance(block, Paragraph)
                and block.style is not None
                and getattr(block.style, "name", None) == heading_style
            ):
                section_indices.append(i)

        section_indices.append(len(blocks))
        total = len(section_indices) - 1

        if total <= 0:
            return 0

        # Generate outputs
        for i in range(total):
            start, end = section_indices[i], section_indices[i + 1]
            section_blocks = blocks[start:end]

            fb = section_blocks[0]
            if isinstance(fb, Paragraph) and fb.style and fb.style.name == heading_style:
                title = fb.text.strip() or f"Section_{i+1}"
            else:
                title = f"Section_{i+1}"

            safe_title = "".join(c for c in title if c.isalnum() or c in " _-").strip()[:50] or f"Section_{i+1}"

            new_doc = Document()

            # Copy page setup
            try:
                s1 = doc.sections[0]
                s2 = new_doc.sections[-1]
                s2.page_width = s1.page_width
                s2.page_height = s1.page_height
                s2.left_margin = s1.left_margin
                s2.right_margin = s1.right_margin
                s2.top_margin = s1.top_margin
                s2.bottom_margin = s1.bottom_margin
            except Exception:
                pass

            # Copy blocks
            for block in section_blocks:
                if isinstance(block, Paragraph):
                    copy_paragraph(block, new_doc, temp_img_dir)
                elif isinstance(block, Table):
                    copy_table(block, new_doc, temp_img_dir)

            out_path = os.path.join(output_dir, f"{i+1:02d}_{safe_title}.docx")

            try:
                new_doc.save(out_path)
            except Exception:
                traceback.print_exc()

        return total

    finally:
        try:
            shutil.rmtree(temp_img_dir)
        except Exception:
            pass

def split_docx_by_heading(input_path, output_dir, heading_style='Heading 1'):
    """
    Splits the provided docx file into multiple docx files based on paragraphs that have
    the specified heading_style (e.g., 'Heading 1'). Returns number of output files.
    """
    os.makedirs(output_dir, exist_ok=True)
    temp_img_dir = os.path.join(output_dir, "temp_images")
    os.makedirs(temp_img_dir, exist_ok=True)

    try:
        doc = Document(input_path)
        blocks = list(iter_block_items(doc))
        if not blocks:
            return 0

        # Determine if the very first block is a heading of desired style safely
        first_block = blocks[0]
        first_is_heading = (
            isinstance(first_block, Paragraph)
            and getattr(first_block, "style", None) is not None
            and getattr(first_block.style, "name", None) == heading_style
        )

        section_indices = []
        if not first_is_heading:
            section_indices.append(0)

        for i, block in enumerate(blocks):
            if (
                isinstance(block, Paragraph)
                and getattr(block, "style", None) is not None
                and getattr(block.style, "name", None) == heading_style
            ):
                section_indices.append(i)

        section_indices.append(len(blocks))
        total_sections = len(section_indices) - 1
        if total_sections <= 0:
            return 0

        for i in range(total_sections):
            start, end = section_indices[i], section_indices[i + 1]
            section_blocks = blocks[start:end]
            if not section_blocks:
                continue

            first_block = section_blocks[0]
            if (
                isinstance(first_block, Paragraph)
                and getattr(first_block, "style", None) is not None
                and getattr(first_block.style, "name", None) == heading_style
            ):
                title = first_block.text.strip()
            else:
                title = "Introduction"

            if not title:
                title = f"Section_{i+1}"

            # Create new document and copy page settings
            current_doc = Document()
            try:
                for section in doc.sections:
                    new_section = current_doc.sections[-1]
                    new_section.page_height = section.page_height
                    new_section.page_width = section.page_width
                    new_section.left_margin = section.left_margin
                    new_section.right_margin = section.right_margin
                    new_section.top_margin = section.top_margin
                    new_section.bottom_margin = section.bottom_margin
                    break
            except Exception:
                pass

            # Copy blocks (paragraphs & tables) for this section
            for block in section_blocks:
                try:
                    if isinstance(block, Paragraph):
                        copy_paragraph(block, current_doc, temp_img_dir)
                    elif isinstance(block, Table):
                        copy_table(block, current_doc, temp_img_dir)
                except Exception:
                    traceback.print_exc()
                    continue

            # Safe filename
            safe_title = "".join(c for c in title if c.isalnum() or c in " _-").strip()[:50]
            if not safe_title:
                safe_title = f"Section_{i+1}"
            out_path = os.path.join(output_dir, f"{i+1:02d}_{safe_title}.docx")

            try:
                current_doc.save(out_path)
            except Exception:
                traceback.print_exc()
                continue

        return total_sections
    finally:
        try:
            if os.path.exists(temp_img_dir):
                shutil.rmtree(temp_img_dir)
        except Exception:
            pass

# ============================================================================
# 4. File Renamer
# ============================================================================

def rename_files_batch(folder_path, old_text, new_text):
    """Batch rename files in a folder."""
    count = 0
    for filename in os.listdir(folder_path):
        if old_text in filename:
            new_filename = filename.replace(old_text, new_text)
            os.rename(
                os.path.join(folder_path, filename),
                os.path.join(folder_path, new_filename)
            )
            count += 1
    return count

def generate_rename_preview_from_excel(excel_path, docx_folder):
    """
    Generate a preview of file renames based on Excel mapping.
    Returns a dictionary with preview list and excel_data.
    """
    try:
        df = pd.read_excel(excel_path)
    except Exception as e:
        return {'error': f'Failed to read Excel: {str(e)}'}
    
    # Find columns (case-insensitive)
    cols_lower = [c.lower().strip() for c in df.columns]
    
    def find_col(names):
        for n in names:
            if n.lower() in cols_lower:
                return df.columns[cols_lower.index(n.lower())]
        return None
    
    doc_col = find_col(["Doc_Name", "Doc Name", "doc_name", "doc name", "DocName", "filename", "File", "FileName"])
    dmc_col = find_col(["DMC_Code", "DMC Code", "dmc_code", "dmc code", "DMC"])
    
    if not doc_col or not dmc_col:
        return {'error': 'Excel must contain Doc Name and DMC Code columns'}
    
    # Extract Excel data for display
    excel_data = []
    for _, row in df.iterrows():
        doc_name = str(row[doc_col]).strip()
        dmc_code = str(row[dmc_col]).strip()
        if doc_name and dmc_code and doc_name.lower() != 'nan' and dmc_code.lower() != 'nan':
            excel_data.append({
                'doc_name': doc_name,
                'dmc_code': dmc_code
            })
    
    # Build mapping
    df[doc_col] = df[doc_col].astype(str).str.strip()
    df[dmc_col] = df[dmc_col].astype(str).str.strip()
    
    mapping = {}
    for _, row in df.iterrows():
        doc_name = str(row[doc_col]).strip()
        dmc_code = str(row[dmc_col]).strip()
        if doc_name and dmc_code and doc_name.lower() != 'nan' and dmc_code.lower() != 'nan':
            mapping[doc_name.lower()] = dmc_code
    
    # Generate preview
    preview = []
    for filename in sorted(os.listdir(docx_folder)):
        if filename.lower().endswith('.docx'):
            base_name = os.path.splitext(filename)[0]
            key = base_name.lower()
            
            if key in mapping:
                dmc_code = mapping[key]
                # Sanitize DMC code for filename
                safe_dmc = "".join(c for c in dmc_code if c.isalnum() or c in (" ", "-", "_")).strip()
                new_filename = f"{safe_dmc}.docx"
                
                # Check if new name already exists
                new_path = os.path.join(docx_folder, new_filename)
                exists = os.path.exists(new_path) and new_path != os.path.join(docx_folder, filename)
                
                preview.append({
                    'original_name': filename,
                    'new_name': new_filename,
                    'dmc_code': dmc_code,
                    'exists': exists,
                    'status': '✗ conflict' if exists else '✓ ready'
                })
            else:
                preview.append({
                    'original_name': filename,
                    'new_name': filename,
                    'dmc_code': '',
                    'exists': False,
                    'status': '⚠ no_mapping'
                })
    
    return {
        'preview': preview,
        'excel_data': excel_data
    }

def execute_rename_from_excel(excel_path, docx_folder, preview_data):
    """
    Execute the rename based on preview data.
    Returns count of renamed files and any errors.
    """
    renamed = 0
    errors = []
    
    for item in preview_data:
        if '✓' in item['status']:  # Check for ready status (with ✓ icon)
            old_path = os.path.join(docx_folder, item['original_name'])
            new_path = os.path.join(docx_folder, item['new_name'])
            
            # Skip if already same name
            if old_path == new_path:
                continue
            
            # Handle conflicts by adding number suffix
            if os.path.exists(new_path):
                base, ext = os.path.splitext(item['new_name'])
                i = 1
                while os.path.exists(os.path.join(docx_folder, f"{base}_{i}{ext}")):
                    i += 1
                new_path = os.path.join(docx_folder, f"{base}_{i}{ext}")
            
            try:
                shutil.move(old_path, new_path)
                renamed += 1
            except Exception as e:
                errors.append(f"Failed to rename {item['original_name']}: {str(e)}")
    
    return renamed, errors

# ============================================================================
# 5. ICN Extractor
# ============================================================================

def extract_icn_from_docx(input_dir, output_dir):
    """Extract images with ICN tags from DOCX files."""
    os.makedirs(output_dir, exist_ok=True)
    
    for filename in os.listdir(input_dir):
        if not filename.lower().endswith('.docx') or filename.startswith('~'):
            continue
        
        docx_path = os.path.join(input_dir, filename)
        base_name = os.path.splitext(filename)[0]
        doc_output_dir = os.path.join(output_dir, base_name)
        
        with zipfile.ZipFile(docx_path, 'r') as docx:
            media_files = sorted([f for f in docx.namelist() if f.startswith('word/media/')])
            if not media_files:
                continue
            
            try:
                xml_content = docx.read("word/document.xml")
                plain_text = ""
                try:
                    root = ET.fromstring(xml_content)
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    for t in root.findall('.//w:t', ns):
                        if t.text:
                            plain_text += t.text
                except ET.ParseError:
                    plain_text = xml_content.decode('utf-8', errors='ignore')
                
                icn_matches = re.findall(r'ICN-\s*([\w\-.]+)', plain_text)
                icn_labels = [f"ICN-{match}" for match in icn_matches]
                
                if not os.path.exists(doc_output_dir):
                    os.makedirs(doc_output_dir)
                
                for i, media_file in enumerate(media_files):
                    if i < len(icn_labels):
                        label = icn_labels[i]
                    else:
                        label = f"image_{i + 1}"
                    
                    ext = os.path.splitext(media_file)[1]
                    safe_label = re.sub(r'[<>:"/\\|?*]', '_', label)
                    out_path = os.path.join(doc_output_dir, f"{safe_label}{ext}")
                    
                    image_data = docx.read(media_file)
                    with open(out_path, "wb") as out_file:
                        out_file.write(image_data)
            except Exception as e:
                print(f"Error processing {filename}: {e}")
                continue

# ============================================================================
# 6. ICN Maker/Generator
# ============================================================================

def generate_icn_code(dmc_code, kpc, xyz, sq, icv, issue, sec):
    """Generate ICN code from DMC code."""
    parts = dmc_code.split("-")
    up_to_unit = f"ICN-{parts[1]}-{parts[2]}-".lstrip('DMC-') + "".join(parts[3:-4]) + f"-{kpc}-{xyz}-{sq}-{icv}-{issue}-{sec}"
    return up_to_unit

def paragraph_has_image(paragraph):
    """Check if a paragraph contains an image."""
    xml = paragraph._element.xml
    return ("<w:drawing" in xml) or ("<w:pict" in xml) or ("<wp:inline" in xml)

def generate_icn_labels(input_dir, output_dir, params):
    """Generate ICN labels for images in DOCX files."""
    os.makedirs(output_dir, exist_ok=True)
    
    current_sq = int(params['sq_start'])
    pad_len = len(params['sq_start'])
    
    for filename in os.listdir(input_dir):
        if not filename.lower().endswith('.docx') or filename.startswith('~'):
            continue
        
        input_path = os.path.join(input_dir, filename)
        dmc_code = os.path.splitext(filename)[0]
        
        doc = Document(input_path)
        paragraphs = doc.paragraphs[:]
        i = 0
        
        while i < len(paragraphs):
            para = paragraphs[i]
            has_image = paragraph_has_image(para)
            
            if has_image:
                icn = generate_icn_code(
                    dmc_code, params['kpc'], params['xyz'],
                    str(current_sq).zfill(pad_len), params['icv'],
                    params['issue'], params['sec']
                )
                current_sq += 1
                if icn:
                    para._p.addnext(doc.add_paragraph(icn)._p)
                    paragraphs = doc.paragraphs[:]
                    i += 1
            i += 1
        
        output_path = os.path.join(output_dir, filename)
        doc.save(output_path)

# ============================================================================
# 7. ICN Validator
# ============================================================================

def validate_adoc_images(adoc_dir, images_dir):
    """Validate ADOC image references against actual image files."""
    IMAGE_PATTERN = re.compile(r'image:?:?(.+?)\[', re.IGNORECASE)
    results = []
    
    for adoc_file in os.listdir(adoc_dir):
        if not adoc_file.endswith('.adoc'):
            continue
        
        adoc_path = os.path.join(adoc_dir, adoc_file)
        dmc_name = os.path.splitext(adoc_file)[0]
        image_folder_path = os.path.join(images_dir, dmc_name)
        
        # Extract referenced images from ADOC
        referenced_images = set()
        try:
            with open(adoc_path, 'r', encoding='utf-8') as f:
                content = f.read()
                matches = IMAGE_PATTERN.findall(content)
                for match in matches:
                    referenced_images.add(match.strip())
        except Exception as e:
            results.append({
                'file': adoc_file,
                'error': str(e),
                'missing': [],
                'unused': []
            })
            continue
        
        # List existing images
        existing_images = set()
        if os.path.isdir(image_folder_path):
            for root, _, files in os.walk(image_folder_path):
                for file in files:
                    if not file.startswith('.'):
                        relative_path = os.path.join(root, file)
                        relative_path = os.path.relpath(relative_path, image_folder_path)
                        existing_images.add(relative_path.replace(os.path.sep, '/'))
        
        # Perform checks
        missing_images = list(referenced_images - existing_images)
        unused_images = list(existing_images - referenced_images)
        
        results.append({
            'file': adoc_file,
            'missing': missing_images,
            'unused': unused_images,
            'status': 'ok' if not missing_images and not unused_images else 'warning'
        })
    
    return results

# ============================================================================
# 8. AsciiDoc to S1000D XML Converter
# ============================================================================

def convert_adoc_to_s1000d(input_path, output_path, ruby_backend_path):
    """
    Convert an AsciiDoc file to S1000D XML format using asciidoctor with Ruby backend.
    
    Args:
        input_path: Path to the .adoc file
        output_path: Path where the .xml file should be saved
        ruby_backend_path: Path to the s1000d1.rb backend file
        
    Returns:
        Tuple of (success: bool, message: str)
    """
    import logging
    import datetime
    
    # Create log file in temp directory
    log_file = os.path.join(tempfile.gettempdir(), 'asciidoctor_conversion.log')
    
    try:
        # Setup logging
        with open(log_file, 'a', encoding='utf-8') as log:
            log.write(f"\n{'='*80}\n")
            log.write(f"Conversion attempt: {datetime.datetime.now()}\n")
            log.write(f"{'='*80}\n")
            
            # Verify input file exists
            if not os.path.exists(input_path):
                error_msg = f"Input file not found: {input_path}"
                log.write(f"ERROR: {error_msg}\n")
                return False, error_msg
            
            log.write(f"Input file: {input_path} (exists: {os.path.exists(input_path)})\n")
            log.write(f"Output file: {output_path}\n")
            
            # Verify Ruby backend exists
            if not os.path.exists(ruby_backend_path):
                error_msg = f"Ruby backend file not found: {ruby_backend_path}"
                log.write(f"ERROR: {error_msg}\n")
                return False, error_msg
            
            log.write(f"Ruby backend: {ruby_backend_path} (exists: {os.path.exists(ruby_backend_path)})\n")
            
            # Get full path to asciidoctor executable
            asciidoctor_exe = get_tool_path('asciidoctor')
            log.write(f"AsciiDoctor path: {asciidoctor_exe}\n")
            log.write(f"AsciiDoctor exists: {os.path.exists(asciidoctor_exe)}\n")
            
            # Build asciidoctor command - use absolute path for Ruby backend
            ruby_backend_abs = os.path.abspath(ruby_backend_path)
            
            cmd = [
                asciidoctor_exe,
                '-r', ruby_backend_abs,
                '-b', 's1000d',
                '-a', 'allow-uri-read',
                '-a', 'source-highlighter=none',
                '--trace',
                '-o', str(output_path),
                str(input_path)
            ]
            
            log.write(f"\nCommand to execute:\n")
            log.write(f"  {' '.join(cmd)}\n")
            log.write(f"\nCommand as list:\n")
            for i, arg in enumerate(cmd):
                log.write(f"  [{i}] {arg}\n")
            
            log.write(f"\nEnvironment:\n")
            log.write(f"  CWD: {os.getcwd()}\n")
            log.write(f"  Python: {sys.executable}\n")
            
            # Run asciidoctor - don't use shell=True to avoid path issues
            log.write(f"\nExecuting subprocess...\n")
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='replace',
                check=True,
                shell=False
            )
            
            log.write(f"\nSubprocess completed successfully\n")
            log.write(f"Return code: {result.returncode}\n")
            
            if result.stdout:
                log.write(f"\nSTDOUT:\n{result.stdout}\n")
            
            if result.stderr:
                log.write(f"\nSTDERR:\n{result.stderr}\n")
            
            # Check if output file was created
            if not os.path.exists(output_path):
                error_msg = "Output file was not created"
                log.write(f"ERROR: {error_msg}\n")
                return False, error_msg
            
            log.write(f"\nOutput file created successfully: {output_path}\n")
            log.write(f"Output file size: {os.path.getsize(output_path)} bytes\n")
            
            success_msg = f"Conversion successful. Log: {log_file}"
            if result.stdout:
                success_msg += f"\n{result.stdout}"
            
            return True, success_msg
        
    except subprocess.CalledProcessError as e:
        with open(log_file, 'a', encoding='utf-8') as log:
            log.write(f"\nERROR: subprocess.CalledProcessError\n")
            log.write(f"Return code: {e.returncode}\n")
            log.write(f"Command: {e.cmd}\n")
            if e.stdout:
                log.write(f"STDOUT:\n{e.stdout}\n")
            if e.stderr:
                log.write(f"STDERR:\n{e.stderr}\n")
        
        error_msg = f"Asciidoctor conversion failed (exit code {e.returncode}). Check log: {log_file}"
        if e.stderr:
            error_msg += f"\nError: {e.stderr[:500]}"
        return False, error_msg
        
    except FileNotFoundError as e:
        with open(log_file, 'a', encoding='utf-8') as log:
            log.write(f"\nERROR: FileNotFoundError\n")
            log.write(f"Error: {str(e)}\n")
            log.write(f"Asciidoctor path attempted: {asciidoctor_exe}\n")
        
        return False, f"Asciidoctor executable not found: {asciidoctor_exe}. Check log: {log_file}"
        
    except Exception as e:
        with open(log_file, 'a', encoding='utf-8') as log:
            log.write(f"\nERROR: Unexpected exception\n")
            log.write(f"Type: {type(e).__name__}\n")
            log.write(f"Error: {str(e)}\n")
            import traceback
            log.write(f"Traceback:\n{traceback.format_exc()}\n")
        
        return False, f"Unexpected error: {str(e)}. Check log: {log_file}"
